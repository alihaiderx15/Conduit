
from __future__ import annotations

from pathlib import Path
import re
from typing import Any

from .common import DependencyUnavailable, FileProcessingError, file_basic_info, safe_output_path
from .models import FileInput, FileKind, ProcessingResult


def extract_text(file: FileInput) -> str:
    path = file.path
    if file.kind in {FileKind.TEXT, FileKind.CODE}:
        return path.read_text(encoding="utf-8", errors="replace")

    if path.suffix.casefold() == ".docx":
        try:
            from docx import Document
        except Exception as exc:
            raise DependencyUnavailable("Word processing requires python-docx.") from exc
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)

    raise FileProcessingError(f"Unsupported document format: {path.suffix}")


def _write_docx_or_text(file: FileInput, text: str, suffix: str) -> Path:
    if file.path.suffix.casefold() == ".docx":
        try:
            from docx import Document
        except Exception as exc:
            raise DependencyUnavailable("Word processing requires python-docx.") from exc
        target = safe_output_path(file.path, suffix, ".docx")
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        doc.save(target)
        return target

    target = safe_output_path(file.path, suffix, file.path.suffix or ".txt")
    target.write_text(text, encoding="utf-8")
    return target


def process(file: FileInput, action: str, params: dict[str, Any]) -> ProcessingResult:
    action = action.casefold().strip()
    text = extract_text(file)

    if action == "inspect":
        data = file_basic_info(file.path)
        words = re.findall(r"\b[\w'-]+\b", text, flags=re.UNICODE)
        data.update({"characters": len(text), "words": len(words), "lines": len(text.splitlines())})
        return ProcessingResult(True, action, f"Inspected {file.path.name}.", file, data=data)

    if action == "extract_text":
        target = safe_output_path(file.path, "extracted_text", ".txt")
        target.write_text(text, encoding="utf-8")
        return ProcessingResult(True, action, "Extracted document text.", file, output_path=target,
                                data={"text": text, "characters": len(text)})

    if action == "word_count":
        words = re.findall(r"\b[\w'-]+\b", text, flags=re.UNICODE)
        return ProcessingResult(True, action, f"Document contains {len(words)} words.", file,
                                data={"words": len(words), "characters": len(text)})

    if action == "reformat":
        # Conservative local cleanup: whitespace only. More subjective formatting
        # can be requested via semantic analysis instead.
        cleaned_lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
        cleaned = "\n".join(line for line in cleaned_lines if line or params.get("keep_blank_lines", False))
        target = _write_docx_or_text(file, cleaned, "reformatted")
        return ProcessingResult(True, action, "Reformatted document whitespace.", file, output_path=target)

    if action == "bullet_points":
        sentences = re.split(r"(?<=[.!?])\s+", text.strip())
        bullets = "\n".join(f"- {s.strip()}" for s in sentences if s.strip())
        target = safe_output_path(file.path, "bullets", ".txt")
        target.write_text(bullets, encoding="utf-8")
        return ProcessingResult(True, action, "Converted document text into bullet points.", file,
                                output_path=target, data={"text": bullets})

    if action == "convert":
        fmt = str(params.get("format", "txt")).casefold().lstrip(".")
        if fmt == "txt":
            target = safe_output_path(file.path, "converted", ".txt")
            target.write_text(text, encoding="utf-8")
        elif fmt == "docx":
            try:
                from docx import Document
            except Exception as exc:
                raise DependencyUnavailable("DOCX conversion requires python-docx.") from exc
            target = safe_output_path(file.path, "converted", ".docx")
            doc = Document()
            for line in text.splitlines():
                doc.add_paragraph(line)
            doc.save(target)
        elif fmt == "pdf":
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import inch
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from xml.sax.saxutils import escape
            except Exception as exc:
                raise DependencyUnavailable("PDF conversion requires reportlab.") from exc

            target = safe_output_path(file.path, "converted", ".pdf")
            styles = getSampleStyleSheet()
            body = styles["BodyText"]
            body.leading = 14
            story = []
            for line in text.splitlines():
                stripped = line.strip()
                if stripped:
                    story.append(Paragraph(escape(stripped), body))
                    story.append(Spacer(1, 0.08 * inch))
                else:
                    story.append(Spacer(1, 0.12 * inch))

            if not story:
                story.append(Paragraph(" ", body))

            pdf = SimpleDocTemplate(
                str(target),
                pagesize=A4,
                rightMargin=0.7 * inch,
                leftMargin=0.7 * inch,
                topMargin=0.7 * inch,
                bottomMargin=0.7 * inch,
                title=file.path.stem,
            )
            pdf.build(story)
        else:
            raise FileProcessingError("Document conversion currently supports txt, docx, and pdf.")
        return ProcessingResult(True, action, f"Converted document to {fmt.upper()}.", file, output_path=target)

    if action in {"summarize", "analyze", "fix", "translate"}:
        defaults = {
            "summarize": "Summarize this document clearly and preserve its important information.",
            "analyze": "Analyze this document, including its key points, structure, and notable issues.",
            "fix": "Correct grammar, spelling, punctuation, and obvious wording errors while preserving meaning.",
            "translate": f"Translate this document into {params.get('language', 'English')} while preserving meaning and structure.",
        }
        return ProcessingResult(
            True, action, f"Prepared document for {action}.", file,
            data={"characters": len(text)}, semantic_text=text,
            semantic_instruction=str(params.get("instruction") or defaults[action]),
        )

    raise FileProcessingError(f"Unsupported document action: {action}")
