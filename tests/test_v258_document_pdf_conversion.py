
from pathlib import Path
from conduit.file_processing import FileProcessingService


def test_txt_to_pdf(tmp_path):
    source = tmp_path / "notes.txt"
    source.write_text("Hello from Conduit.\nThis is PDF conversion.", encoding="utf-8")
    service = FileProcessingService(state_path=tmp_path/"state.json")
    service.register_dropped_file(source)
    result = service.process(action="convert", parameters={"format": "pdf"})
    assert result.success
    assert result.output_path is not None
    assert result.output_path.suffix == ".pdf"
    assert result.output_path.exists()
    assert result.output_path.stat().st_size > 0


def test_docx_to_pdf(tmp_path):
    from docx import Document
    source = tmp_path / "notes.docx"
    doc = Document()
    doc.add_paragraph("Hello from DOCX.")
    doc.add_paragraph("Convert me to PDF.")
    doc.save(source)

    service = FileProcessingService(state_path=tmp_path/"state.json")
    service.register_dropped_file(source)
    result = service.process(action="convert", parameters={"format": "pdf"})
    assert result.success
    assert result.output_path is not None
    assert result.output_path.suffix == ".pdf"
    assert result.output_path.exists()
    assert result.output_path.stat().st_size > 0


def test_version_258():
    root = Path(__file__).resolve().parents[1]
    assert 'version = "3.1.8"' in (root/"pyproject.toml").read_text(encoding="utf-8")
    assert '"reportlab>=4.2.0"' in (root/"pyproject.toml").read_text(encoding="utf-8")
